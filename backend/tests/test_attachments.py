"""Attachments：parsers / 上傳 API / 刪除 / PRD prompt 注入。"""
from __future__ import annotations

import io
import shutil
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

import app as appmod
from parsers import parse


# ============================================================
#  Parsers
# ============================================================
def test_parse_text_md(tmp_path: Path):
    p = tmp_path / "spec.md"
    p.write_text("# Hello\n\nFR-1 訪客結帳", encoding="utf-8")
    text, err = parse(p, "text/markdown", "spec.md")
    assert err == ""
    assert "FR-1" in text and "訪客結帳" in text


def test_parse_text_unknown_mime_fallback(tmp_path: Path):
    """未知 mime，但能 UTF-8 解碼 → 仍視為文字。"""
    p = tmp_path / "weird.dat"
    p.write_text("plain text content", encoding="utf-8")
    text, err = parse(p, "application/x-custom", "weird.dat")
    assert "plain text" in text
    assert err == "" or "unsupported" in err  # 容忍兩種行為


def test_parse_image_no_tesseract_fallback(tmp_path: Path, monkeypatch):
    """tesseract 不在 PATH 時，圖片回 fallback message，不 raise。"""
    monkeypatch.setattr("shutil.which", lambda _cmd: None)
    p = tmp_path / "fake.png"
    p.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 32)  # 假 PNG header
    text, err = parse(p, "image/png", "fake.png")
    assert text == ""
    assert "tesseract" in err.lower()


def test_parse_pdf_minimal(tmp_path: Path):
    """用 pypdf 寫一個極簡 PDF（無文字）→ 不 raise。"""
    try:
        from pypdf import PdfWriter
    except ImportError:
        pytest.skip("pypdf not installed")
    p = tmp_path / "blank.pdf"
    w = PdfWriter()
    w.add_blank_page(width=72, height=72)
    with open(p, "wb") as f:
        w.write(f)
    text, err = parse(p, "application/pdf", "blank.pdf")
    assert err == ""  # 解析成功（即使無文字）


def test_parse_docx_minimal(tmp_path: Path):
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")
    p = tmp_path / "spec.docx"
    doc = Document()
    doc.add_paragraph("產品需求：訪客結帳 + Apple Pay")
    doc.add_paragraph("並發：5,000")
    doc.save(p)
    text, err = parse(p, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "spec.docx")
    assert err == ""
    assert "Apple Pay" in text
    assert "5,000" in text


# ============================================================
#  Attachments API
# ============================================================
def test_attachments_upload_list_download_delete(tmp_db):
    with TestClient(appmod.app) as c:
        tid = c.post("/api/projects", json={"name": "t"}).json()["thread_id"]
        files = {"file": ("req.md",
                          io.BytesIO("# 需求\n\nFR-1 訪客結帳".encode("utf-8")),
                          "text/markdown")}
        r = c.post(f"/api/stage/prd/{tid}/attachments", files=files)
        assert r.status_code == 200
        body = r.json()
        assert body["has_parsed_text"] is True
        assert body["parse_error"] in (None, "")
        fid = body["file_id"]

        # list
        lst = c.get(f"/api/stage/prd/{tid}/attachments").json()["attachments"]
        assert len(lst) == 1 and lst[0]["file_id"] == fid

        # download
        dl = c.get(f"/api/stage/prd/{tid}/attachments/{fid}/content")
        assert dl.status_code == 200
        assert "訪客結帳".encode("utf-8") in dl.content

        # delete
        dr = c.delete(f"/api/stage/prd/{tid}/attachments/{fid}")
        assert dr.status_code == 200 and dr.json()["deleted"] == fid

        # list empty
        assert c.get(f"/api/stage/prd/{tid}/attachments").json()["attachments"] == []


def test_upload_to_unknown_thread_404(tmp_db):
    with TestClient(appmod.app) as c:
        files = {"file": ("x.md", b"hi", "text/markdown")}
        r = c.post("/api/stage/prd/nonexistent_thread/attachments", files=files)
        assert r.status_code == 404
        assert r.json()["detail"]["category"] == "thread_not_found"


def test_delete_unknown_attachment_404(tmp_db):
    with TestClient(appmod.app) as c:
        tid = c.post("/api/projects", json={"name": "t"}).json()["thread_id"]
        r = c.delete(f"/api/stage/prd/{tid}/attachments/never_existed")
        assert r.status_code == 404


# ============================================================
#  PRD prompt injection（M1.3 path-passing：列 abs_path + READ 指令）
# ============================================================
def test_prd_prompt_passes_attachment_path(tmp_db):
    """上傳 md 附件 → generate PRD → prompt 含絕對路徑與 READ 指令（path-passing）。"""
    from plugin_api import ModelAdapter
    import plugin_loader as L
    from persistence import dal
    from workflow_engine import WorkflowEngine

    reg = L.load_all()
    captured: dict = {}
    def fake_invoke(prompt: str) -> str:
        captured["prompt"] = prompt
        return ("# Product Requirements Document\n"
                "## 1. Overview\nDerived from attachment.\n"
                "## 3. Functional Requirements\n- `FR-1`: 訪客結帳\n"
                "## 4. Non-Functional Requirements\n- `NFR-1`: 5,000 並發\n"
                "[PRD_READY]")
    reg.model_adapters["claude-cli"] = ModelAdapter(
        model_choice="claude-cli", invoke=fake_invoke, is_available=lambda: True,
        description="mock", max_context_tokens=1000,
        prompt_budget_tokens=900, response_budget_tokens=100,
    )

    dal.create_project("t1", "test")
    uploads = dal.uploads_dir() / "t1"
    uploads.mkdir(parents=True, exist_ok=True)
    rel = "t1/abc123.md"
    (dal.uploads_dir() / rel).write_text("# 競品調研\n\nApple Pay 對 iOS 轉換率關鍵",
                                          encoding="utf-8")
    dal.add_attachment(
        file_id="abc123", thread_id="t1", stage_id="prd",
        filename="research.md", mime="text/markdown", size_bytes=64,
        content_path=rel,
        parsed_text=None,  # M1.3：parser 不再強制 inline；handler 走 path
    )

    engine = WorkflowEngine(reg)
    engine.dispatch(thread_id="t1", stage_id="prd", op="generate")

    prompt = captured["prompt"]
    # 1. READ 指令在
    assert "READ" in prompt and "Read tool" in prompt
    # 2. 絕對路徑在
    expected_abs = str(dal.uploads_dir() / rel)
    assert expected_abs in prompt
    # 3. 原始檔名 / mime 也標示
    assert "research.md" in prompt
    assert "text/markdown" in prompt
    # 4. 不再 inline parsed_text（path-passing 應該不含內容）
    assert "Apple Pay" not in prompt
    assert "<<< attachment:" not in prompt


def test_prd_prompt_without_attachments_has_marker(tmp_db):
    """沒附件時 prompt 仍含 attachments section（顯示 "no attached files"），不 raise。"""
    from plugin_api import ModelAdapter
    import plugin_loader as L
    from persistence import dal
    from workflow_engine import WorkflowEngine

    reg = L.load_all()
    captured: dict = {}
    reg.model_adapters["claude-cli"] = ModelAdapter(
        model_choice="claude-cli",
        invoke=lambda p: (captured.update(prompt=p),
                          "# PRD\n## 1. Overview\nx\n## 3. Functional Requirements\n- `FR-1`: y\n## 4. Non-Functional Requirements\n- `NFR-1`: z\n[PRD_READY]")[1],
        is_available=lambda: True,
        description="mock", max_context_tokens=1000,
        prompt_budget_tokens=900, response_budget_tokens=100,
    )
    dal.create_project("t2", "test")
    engine = WorkflowEngine(reg)
    engine.dispatch(thread_id="t2", stage_id="prd", op="generate")
    assert "no attached files" in captured["prompt"].lower()


# ============================================================
#  _format_attachments unit tests（M1.3 path-passing vs inline fallback）
# ============================================================
def test_format_attachments_path_list():
    """所有 attachment 都有 abs_path → path-list 格式 + READ 指令，無 inline 內容。"""
    from plugins.builtin_core_stages.prd_stage import _format_attachments

    out = _format_attachments([
        {"abs_path": "/srv/uploads/t1/aaa.png", "filename": "screen.png",
         "mime": "image/png", "size_bytes": 12345,
         "parsed_text": "[OCR text that should NOT appear]"},
        {"abs_path": "/srv/uploads/t1/bbb.pdf", "filename": "spec.pdf",
         "mime": "application/pdf", "size_bytes": 99999,
         "parsed_text": "[PDF text that should NOT appear]"},
    ])

    # READ 指令在
    assert "READ" in out and "Read tool" in out
    # 兩條 path 在
    assert "/srv/uploads/t1/aaa.png" in out
    assert "/srv/uploads/t1/bbb.pdf" in out
    # 原始檔名 / mime / size 在
    assert "screen.png" in out and "image/png" in out and "12345" in out
    assert "spec.pdf" in out and "application/pdf" in out
    # parsed_text 不應該被列出（path-passing 模式）
    assert "OCR text" not in out
    assert "PDF text" not in out
    # 不應該有 inline marker
    assert "<<< attachment:" not in out


def test_format_attachments_inline_fallback():
    """缺 abs_path 時退回 inline parsed_text + marker（保留 M1.1 行為）。"""
    from plugins.builtin_core_stages.prd_stage import _format_attachments

    out = _format_attachments([
        {"filename": "notes.txt", "mime": "text/plain", "size_bytes": 42,
         "parsed_text": "Apple Pay 對 iOS 轉換率關鍵"},
    ])

    assert "<<< attachment: notes.txt" in out
    assert "Apple Pay" in out
    assert "<<< end of notes.txt >>>" in out
    # 不該出現 path-list 模式的 READ 指令
    assert "Read tool" not in out


def test_format_attachments_mixed_falls_back():
    """有一筆缺 abs_path → 整批走 inline fallback（保守策略）。"""
    from plugins.builtin_core_stages.prd_stage import _format_attachments

    out = _format_attachments([
        {"abs_path": "/srv/uploads/t1/a.md", "filename": "a.md",
         "mime": "text/markdown", "size_bytes": 10, "parsed_text": "AAA content"},
        {"filename": "b.md", "mime": "text/markdown",
         "size_bytes": 10, "parsed_text": "BBB content"},
    ])
    assert "<<< attachment: a.md" in out
    assert "AAA content" in out
    assert "BBB content" in out
    assert "Read tool" not in out


def test_format_attachments_empty():
    from plugins.builtin_core_stages.prd_stage import _format_attachments
    assert _format_attachments([]) == "(no attached files)"
