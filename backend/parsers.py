"""附件解析：把上傳檔變成純文字 inline 進 SA prompt。

對應檔案類型：
- text / md / csv / json / xml / yaml / html / log → 直接讀
- PDF → pypdf
- DOCX → python-docx
- 圖片 (png/jpg/webp/gif/bmp) → pytesseract OCR（tesseract 未安裝 → fallback message）

純文字內容上限 80KB（防 prompt 爆），超出截斷。
解析失敗都回 (text="", error="<原因>")，不 raise；caller 自行決定要不要 inline 進 prompt。
"""
from __future__ import annotations

import logging
import shutil
from pathlib import Path
from typing import Tuple

log = logging.getLogger("parsers")

# ---------- 類型對應 ----------
_TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".xml",
    ".yaml", ".yml", ".html", ".htm", ".log", ".rst", ".ini", ".toml",
}
_TEXT_MIMES = {
    "text/plain", "text/markdown", "text/csv", "text/tab-separated-values",
    "application/json", "application/xml", "text/xml",
    "application/x-yaml", "text/yaml", "text/html",
}
_PDF_EXTS = {".pdf"}
_PDF_MIMES = {"application/pdf"}
_DOCX_EXTS = {".docx"}
_DOCX_MIMES = {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}
_IMAGE_MIMES = {"image/png", "image/jpeg", "image/webp", "image/gif", "image/bmp"}

# OCR 語言設定：中文（繁體）+ 英文；tesseract 未裝 chi_tra 時 fallback 至 eng
_OCR_LANGS_PREF = ("chi_tra+eng", "eng")

# 上限與標記
_MAX_TEXT_PER_FILE = 80_000          # ~80KB；超出截斷
_TRUNC_NOTE = "\n\n[…內容過長已截斷…]"


# ============================================================
def parse(path: str | Path, mime: str, filename: str) -> Tuple[str, str]:
    """解析檔案成純文字。

    Returns
    -------
    (parsed_text, parse_error) : tuple[str, str]
        失敗時 parsed_text 為空字串、parse_error 含原因；caller 不需 try/except。
    """
    p = Path(path)
    ext = Path(filename).suffix.lower()

    try:
        if mime in _TEXT_MIMES or ext in _TEXT_EXTS:
            return _parse_text(p), ""
        if mime in _PDF_MIMES or ext in _PDF_EXTS:
            return _parse_pdf(p), ""
        if mime in _DOCX_MIMES or ext in _DOCX_EXTS:
            return _parse_docx(p), ""
        if mime in _IMAGE_MIMES or ext in _IMAGE_EXTS:
            return _parse_image(p)
        # 未知類型：嘗試純文字讀（latin-1 fallback）
        try:
            return _parse_text(p), ""
        except Exception:
            return "", f"unsupported file type: mime={mime!r}, ext={ext!r}"
    except Exception as exc:  # noqa: BLE001
        log.exception("parse failed: %s", filename)
        return "", f"parse error: {type(exc).__name__}: {exc}"


# ============================================================
def _truncate(text: str) -> str:
    return text if len(text) <= _MAX_TEXT_PER_FILE else text[:_MAX_TEXT_PER_FILE] + _TRUNC_NOTE


def _parse_text(path: Path) -> str:
    raw = path.read_bytes()
    try:
        return _truncate(raw.decode("utf-8"))
    except UnicodeDecodeError:
        return _truncate(raw.decode("latin-1", errors="replace"))


def _parse_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001
            pages.append(f"[page {i + 1} extract failed: {exc}]")
    return _truncate("\n\n--- page break ---\n\n".join(pages))


def _parse_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return _truncate("\n".join(parts))


def _parse_image(path: Path) -> Tuple[str, str]:
    """OCR 圖片。tesseract 未安裝 → fallback message。"""
    if shutil.which("tesseract") is None:
        return "", (
            "tesseract 未安裝；圖片 OCR 跳過。啟用：`brew install tesseract tesseract-lang` "
            "後重啟 backend。"
        )
    try:
        import pytesseract
        from PIL import Image
    except ImportError as exc:
        return "", f"OCR 依賴未安裝（{exc}）"
    img = Image.open(path)
    last_exc: Exception | None = None
    for lang in _OCR_LANGS_PREF:
        try:
            text = pytesseract.image_to_string(img, lang=lang)
            return _truncate(text.strip()), ""
        except Exception as exc:  # noqa: BLE001
            last_exc = exc
            continue
    return "", f"OCR error: {type(last_exc).__name__}: {last_exc}"
